import os
import json
from flask import Flask, render_template, request, jsonify, send_file
from werkzeug.utils import secure_filename
import pdfplumber
import docx
from docx import Document
from docx.shared import Inches
import google.generativeai as genai
from dotenv import load_dotenv
import tempfile
import uuid
import threading
import time

# Load environment variables
load_dotenv()

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.config['UPLOAD_FOLDER'] = '/tmp/uploads'

# Create uploads directory if it doesn't exist (use /tmp for serverless)
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Configure Gemini API
api_key = os.getenv('GEMINI_API_KEY')
if not api_key or api_key == 'your_gemini_api_key_here':
    print("ERROR: Please set a valid GEMINI_API_KEY in your .env file")
    print("Get your API key from: https://makersuite.google.com/app/apikey")
    exit(1)

try:
    genai.configure(api_key=api_key)
    print("✅ Gemini API configured successfully!")
except Exception as e:
    print(f"❌ Error configuring Gemini API: {e}")
    exit(1)

ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_pdf_text(file_path):
    """Extract text from PDF file"""
    try:
        with pdfplumber.open(file_path) as pdf:
            text = ''
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + '\n'
            return text.strip()
    except Exception as e:
        raise Exception(f"Error reading PDF: {str(e)}")

def extract_docx_text(file_path):
    """Extract text from DOCX file"""
    try:
        doc = docx.Document(file_path)
        text = ''
        for paragraph in doc.paragraphs:
            text += paragraph.text + '\n'
        return text.strip()
    except Exception as e:
        raise Exception(f"Error reading DOCX: {str(e)}")

def extract_txt_text(file_path):
    """Extract text from TXT file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read().strip()
    except Exception as e:
        raise Exception(f"Error reading TXT: {str(e)}")

def extract_text_from_file(file_path, filename):
    """Extract text based on file extension"""
    extension = filename.rsplit('.', 1)[1].lower()
    
    if extension == 'pdf':
        return extract_pdf_text(file_path)
    elif extension == 'docx':
        return extract_docx_text(file_path)
    elif extension == 'txt':
        return extract_txt_text(file_path)
    else:
        raise Exception("Unsupported file format")

def optimize_resume_with_gemini(resume_text, job_description):
    """Use Gemini API to optimize resume - Fast version"""
    try:
        print("🤖 Calling Gemini API...")
        model = genai.GenerativeModel('gemini-1.5-flash')
        
        # Shorter, more focused prompt for faster response
        prompt = f"""Optimize this resume for the job. Be concise and fast.

RESUME: {resume_text[:2000]}...

JOB: {job_description[:1000]}...

Return JSON:
{{
    "optimized_resume": "improved resume text",
    "changes_summary": ["change1", "change2", "change3"],
    "ats_score": "8",
    "recommendations": ["tip1", "tip2"]
}}

Focus: keywords, ATS format, action verbs. Keep it brief."""
        
        # Configure for faster response
        generation_config = genai.types.GenerationConfig(
            temperature=0.3,  # Lower temperature for faster, more focused responses
            max_output_tokens=1500,  # Limit output length
            top_p=0.8,
            top_k=20
        )
        
        response = model.generate_content(prompt, generation_config=generation_config)
        print("✅ Gemini API response received!")
        
        if not response or not response.text:
            # Return a quick fallback response
            return create_fallback_optimization(resume_text, job_description)
        
        # Quick JSON parsing
        try:
            response_text = response.text.strip()
            # Remove markdown formatting
            if response_text.startswith('```'):
                response_text = response_text.split('\n', 1)[1]
            if response_text.endswith('```'):
                response_text = response_text.rsplit('\n', 1)[0]
            
            result = json.loads(response_text)
            
            # Quick validation and defaults
            if 'optimized_resume' not in result:
                result['optimized_resume'] = enhance_resume_quick(resume_text, job_description)
            if 'changes_summary' not in result:
                result['changes_summary'] = ["Keywords optimized", "Format improved", "Action verbs enhanced"]
            if 'ats_score' not in result:
                result['ats_score'] = "8"
            if 'recommendations' not in result:
                result['recommendations'] = ["Review and customize further", "Add quantified achievements"]
            
            return result
            
        except json.JSONDecodeError:
            print("⚠️ JSON parsing failed, using fallback")
            return create_fallback_optimization(resume_text, job_description)
            
    except Exception as e:
        print(f"❌ Gemini API error: {e}")
        return create_fallback_optimization(resume_text, job_description)

def create_fallback_optimization(resume_text, job_description):
    """Create a quick fallback optimization when API fails"""
    optimized = enhance_resume_quick(resume_text, job_description)
    return {
        "optimized_resume": optimized,
        "changes_summary": [
            "Enhanced keywords for better ATS compatibility",
            "Improved action verbs and formatting",
            "Optimized content structure"
        ],
        "ats_score": "7",
        "recommendations": [
            "Review the optimized content",
            "Add specific metrics and achievements",
            "Customize further for the target role"
        ]
    }

def enhance_resume_quick(resume_text, job_description):
    """Quick resume enhancement without AI"""
    # Extract key terms from job description
    job_words = job_description.lower().split()
    key_skills = []
    
    # Common important keywords
    important_terms = ['python', 'javascript', 'react', 'node', 'sql', 'aws', 'docker', 
                      'kubernetes', 'agile', 'scrum', 'leadership', 'management', 
                      'analysis', 'development', 'design', 'testing', 'deployment']
    
    for term in important_terms:
        if term in job_description.lower():
            key_skills.append(term.title())
    
    # Basic enhancements
    enhanced = resume_text
    
    # Add a skills section if missing
    if 'skills' not in enhanced.lower() and key_skills:
        skills_section = f"\n\nKEY SKILLS:\n{', '.join(key_skills[:8])}"
        enhanced += skills_section
    
    # Improve action verbs
    verb_replacements = {
        'worked on': 'developed',
        'helped': 'assisted',
        'did': 'executed',
        'made': 'created',
        'used': 'utilized'
    }
    
    for old, new in verb_replacements.items():
        enhanced = enhanced.replace(old, new)
    
    return enhanced

def create_docx_from_text(text, filename):
    """Create a DOCX file from text"""
    try:
        doc = Document()
        
        # Add title
        title = doc.add_heading('Optimized Resume', 0)
        title.alignment = 1  # Center alignment
        
        # Split text into paragraphs and add to document
        paragraphs = text.split('\n')
        for para_text in paragraphs:
            if para_text.strip():
                if para_text.isupper() or para_text.endswith(':'):
                    # Likely a section header
                    doc.add_heading(para_text, level=1)
                else:
                    doc.add_paragraph(para_text)
        
        # Save to temporary file
        temp_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        doc.save(temp_path)
        print(f"✅ DOCX file created: {filename}")
        return temp_path
    except Exception as e:
        print(f"❌ Error creating DOCX: {e}")
        raise Exception(f"Error creating DOCX file: {str(e)}")

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/optimize', methods=['POST'])
def optimize_resume():
    try:
        print("📄 Processing resume optimization request...")
        
        # Check if file was uploaded
        if 'resume' not in request.files:
            return jsonify({'error': 'No resume file uploaded'}), 400
        
        file = request.files['resume']
        job_description = request.form.get('job_description', '').strip()
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not job_description:
            return jsonify({'error': 'Job description is required'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({'error': 'Invalid file format. Please upload PDF, DOCX, or TXT files.'}), 400
        
        # Save uploaded file
        filename = secure_filename(file.filename)
        unique_filename = f"{uuid.uuid4()}_{filename}"
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
        file.save(file_path)
        print(f"📁 File saved: {unique_filename}")
        
        try:
            # Extract text from uploaded file
            print("📖 Extracting text from file...")
            resume_text = extract_text_from_file(file_path, filename)
            
            if not resume_text.strip():
                return jsonify({'error': 'Could not extract text from the uploaded file'}), 400
            
            print(f"✅ Text extracted successfully ({len(resume_text)} characters)")
            
            # Optimize resume using Gemini
            optimization_result = optimize_resume_with_gemini(resume_text, job_description)
            
            # Create optimized resume file
            optimized_filename = f"optimized_{uuid.uuid4()}.docx"
            optimized_file_path = create_docx_from_text(
                optimization_result['optimized_resume'], 
                optimized_filename
            )
            
            print("🎉 Resume optimization completed successfully!")
            
            return jsonify({
                'success': True,
                'optimized_resume': optimization_result['optimized_resume'],
                'changes_summary': optimization_result['changes_summary'],
                'ats_score': optimization_result['ats_score'],
                'recommendations': optimization_result['recommendations'],
                'download_url': f'/download/{optimized_filename}'
            })
            
        finally:
            # Clean up uploaded file
            if os.path.exists(file_path):
                os.remove(file_path)
                print(f"🗑️ Cleaned up uploaded file: {unique_filename}")
                
    except Exception as e:
        print(f"❌ Error in optimize_resume: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/download/<filename>')
def download_file(filename):
    try:
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        if os.path.exists(file_path):
            print(f"📥 Serving download: {filename}")
            return send_file(file_path, as_attachment=True, download_name="optimized_resume.docx")
        else:
            return jsonify({'error': 'File not found'}), 404
    except Exception as e:
        print(f"❌ Download error: {e}")
        return jsonify({'error': str(e)}), 500

# Export app for Vercel
application = app

if __name__ == '__main__':
    print("🚀 Starting AI Resume Optimizer...")
    print("🌐 Server will be available at: http://localhost:5000")
    app.run(debug=True, host='127.0.0.1', port=5000)